"""
GENESIS - Document Generator
Genera documentos reales en multiples formatos: MD, HTML, TXT, PDF, DOCX.

Cada formato tiene su renderer especializado con estilo JARVIS:
- TXT: Texto plano con bordes ASCII
- MD: Markdown con headers, listas, bloques
- HTML: Pagina completa con CSS JARVIS (cyan/dark theme)
- PDF: ReportLab con tipografia monospace + colores JARVIS
- DOCX: python-docx con estilos corporativos
"""
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

from core.logger import GenesisLogger


class DocumentGenerator:
    """
    Generador de documentos multi-formato.

    Uso:
        gen = DocumentGenerator()
        path = gen.generate(
            title="Mi Reporte",
            content="Contenido del documento...",
            fmt="pdf",
        )
    """

    SUPPORTED_FORMATS = {"txt", "md", "html", "pdf", "docx"}

    # Directorio por defecto para documentos generados
    DEFAULT_DIR = "generated_docs"

    def __init__(self, output_dir: str = ""):
        self.log = GenesisLogger().get_child("doc_gen")
        base = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.output_dir = Path(output_dir) if output_dir else base / self.DEFAULT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.generated_docs: list[dict] = []

    def generate(self, title: str, content: str, fmt: str = "md",
                 author: str = "GENESIS AI",
                 subtitle: str = "",
                 sections: list[dict] = None,
                 output_path: str = "") -> dict:
        """
        Genera un documento en el formato especificado.

        Args:
            title: Titulo del documento
            content: Contenido principal (texto plano o markdown)
            fmt: Formato de salida (txt, md, html, pdf, docx)
            author: Autor del documento
            subtitle: Subtitulo opcional
            sections: Lista de secciones [{title, content}, ...]
            output_path: Ruta personalizada (opcional)

        Returns:
            dict con 'path', 'format', 'size', 'title' o 'error'
        """
        fmt = fmt.lower().strip().lstrip(".")
        if fmt not in self.SUPPORTED_FORMATS:
            return {
                "error": f"Formato no soportado: {fmt}",
                "supported": list(self.SUPPORTED_FORMATS),
            }

        # Nombre de archivo seguro
        safe_title = self._safe_filename(title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if output_path:
            filepath = Path(output_path)
        else:
            filepath = self.output_dir / f"{safe_title}_{timestamp}.{fmt}"

        filepath.parent.mkdir(parents=True, exist_ok=True)

        try:
            if fmt == "txt":
                self._render_txt(filepath, title, content, author, subtitle, sections)
            elif fmt == "md":
                self._render_md(filepath, title, content, author, subtitle, sections)
            elif fmt == "html":
                self._render_html(filepath, title, content, author, subtitle, sections)
            elif fmt == "pdf":
                self._render_pdf(filepath, title, content, author, subtitle, sections)
            elif fmt == "docx":
                self._render_docx(filepath, title, content, author, subtitle, sections)

            size = filepath.stat().st_size
            doc_info = {
                "path": str(filepath),
                "format": fmt,
                "size": size,
                "size_human": self._format_size(size),
                "title": title,
                "timestamp": time.time(),
            }
            self.generated_docs.append(doc_info)
            self.log.info(f"Documento generado: {filepath} ({self._format_size(size)})")
            return doc_info

        except Exception as e:
            self.log.error(f"Error generando documento: {e}")
            return {"error": str(e)}

    def generate_report(self, genesis=None) -> str:
        """Genera reporte del sistema de documentos."""
        lines = [
            "  === DOCUMENT GENERATOR ===",
            f"  Directorio: {self.output_dir}",
            f"  Formatos: {', '.join(sorted(self.SUPPORTED_FORMATS))}",
            f"  Documentos generados: {len(self.generated_docs)}",
        ]
        if self.generated_docs:
            lines.append("")
            lines.append("  Ultimos documentos:")
            for doc in self.generated_docs[-5:]:
                lines.append(
                    f"    [{doc['format'].upper()}] {doc['title']} "
                    f"({doc['size_human']})"
                )
        return "\n".join(lines)

    def get_stats(self) -> dict:
        """Estadisticas del generador."""
        fmt_counts = {}
        for doc in self.generated_docs:
            f = doc.get("format", "?")
            fmt_counts[f] = fmt_counts.get(f, 0) + 1

        return {
            "total_generated": len(self.generated_docs),
            "by_format": fmt_counts,
            "output_dir": str(self.output_dir),
        }

    def save(self):
        """Placeholder para compatibilidad con save_all()."""
        pass

    def clear(self):
        """Limpia historial de documentos generados."""
        self.generated_docs.clear()

    def status(self) -> str:
        """Estado breve."""
        return f"  DocGen: {len(self.generated_docs)} docs | Dir: {self.output_dir}"

    # ================================================================
    # RENDERERS
    # ================================================================

    def _render_txt(self, filepath: Path, title: str, content: str,
                    author: str, subtitle: str, sections: list):
        """Genera documento TXT con formato ASCII."""
        width = 70
        lines = []
        lines.append("=" * width)
        lines.append(title.upper().center(width))
        if subtitle:
            lines.append(subtitle.center(width))
        lines.append(f"Autor: {author}  |  Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}".center(width))
        lines.append("=" * width)
        lines.append("")

        if content:
            lines.append(content)
            lines.append("")

        if sections:
            for sec in sections:
                lines.append("-" * width)
                lines.append(f"  {sec.get('title', 'Sin titulo').upper()}")
                lines.append("-" * width)
                lines.append(sec.get("content", ""))
                lines.append("")

        lines.append("=" * width)
        lines.append(f"Generado por {author} — {datetime.now().isoformat()}")
        lines.append("=" * width)

        filepath.write_text("\n".join(lines), encoding="utf-8")

    def _render_md(self, filepath: Path, title: str, content: str,
                   author: str, subtitle: str, sections: list):
        """Genera documento Markdown."""
        lines = []
        lines.append(f"# {title}")
        if subtitle:
            lines.append(f"*{subtitle}*")
        lines.append("")
        lines.append(f"> **Autor:** {author}  ")
        lines.append(f"> **Fecha:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  ")
        lines.append("")
        lines.append("---")
        lines.append("")

        if content:
            lines.append(content)
            lines.append("")

        if sections:
            for sec in sections:
                lines.append(f"## {sec.get('title', 'Sin titulo')}")
                lines.append("")
                lines.append(sec.get("content", ""))
                lines.append("")

        lines.append("---")
        lines.append(f"*Generado por {author} — {datetime.now().isoformat()}*")

        filepath.write_text("\n".join(lines), encoding="utf-8")

    def _render_html(self, filepath: Path, title: str, content: str,
                     author: str, subtitle: str, sections: list):
        """Genera documento HTML con estilo JARVIS."""
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

        sections_html = ""
        if sections:
            for sec in sections:
                sec_title = sec.get("title", "")
                sec_content = sec.get("content", "").replace("\n", "<br>")
                sections_html += f"""
                <div class="section">
                    <h2>{sec_title}</h2>
                    <div class="section-content">{sec_content}</div>
                </div>"""

        content_html = content.replace("\n", "<br>") if content else ""

        html = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} — GENESIS</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700;900&family=Share+Tech+Mono&display=swap');

        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Share Tech Mono', 'Consolas', monospace;
            background: #020a12;
            color: #c8dce8;
            padding: 40px;
            line-height: 1.8;
        }}
        .document {{
            max-width: 800px;
            margin: 0 auto;
            border: 1px solid rgba(0,180,255,0.2);
            padding: 40px;
            background: rgba(0,180,255,0.02);
        }}
        .header {{
            border-bottom: 2px solid #00b4ff;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        h1 {{
            font-family: 'Orbitron', sans-serif;
            color: #00b4ff;
            font-size: 1.8em;
            letter-spacing: 4px;
            text-shadow: 0 0 20px rgba(0,180,255,0.3);
        }}
        .subtitle {{
            color: #5a8aaa;
            font-size: 1em;
            margin-top: 4px;
        }}
        .meta {{
            color: #2a4a5a;
            font-size: 0.8em;
            margin-top: 10px;
            letter-spacing: 1px;
        }}
        .content {{
            margin: 20px 0;
            white-space: pre-wrap;
        }}
        .section {{
            margin: 24px 0;
            border-left: 3px solid rgba(0,180,255,0.3);
            padding-left: 16px;
        }}
        h2 {{
            font-family: 'Orbitron', sans-serif;
            color: #00e5ff;
            font-size: 1.1em;
            letter-spacing: 2px;
            margin-bottom: 10px;
        }}
        .section-content {{
            color: #c8dce8;
        }}
        .footer {{
            border-top: 1px solid rgba(0,180,255,0.15);
            padding-top: 16px;
            margin-top: 30px;
            color: #2a4a5a;
            font-size: 0.75em;
            text-align: center;
            letter-spacing: 1px;
        }}
        code {{
            background: rgba(0,180,255,0.08);
            border: 1px solid rgba(0,180,255,0.15);
            padding: 1px 5px;
            border-radius: 2px;
            color: #00e5ff;
        }}
        @media print {{
            body {{ background: white; color: #1a1a2e; padding: 20px; }}
            .document {{ border: none; background: none; }}
            h1 {{ color: #0066cc; text-shadow: none; }}
            h2 {{ color: #0088cc; }}
            .meta, .footer {{ color: #888; }}
            .section {{ border-left-color: #0066cc; }}
        }}
    </style>
</head>
<body>
    <div class="document">
        <div class="header">
            <h1>{title}</h1>
            {"<div class='subtitle'>" + subtitle + "</div>" if subtitle else ""}
            <div class="meta">Autor: {author} | Fecha: {date_str}</div>
        </div>
        <div class="content">{content_html}</div>
        {sections_html}
        <div class="footer">
            Generado por {author} // {datetime.now().isoformat()} // GENESIS Document System
        </div>
    </div>
</body>
</html>"""
        filepath.write_text(html, encoding="utf-8")

    def _render_pdf(self, filepath: Path, title: str, content: str,
                    author: str, subtitle: str, sections: list):
        """Genera documento PDF con ReportLab (estilo JARVIS)."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.lib.colors import HexColor
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, HRFlowable
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError:
            # Fallback: generar HTML y renombrar como referencia
            html_path = filepath.with_suffix(".html")
            self._render_html(html_path, title, content, author, subtitle, sections)
            html_path.rename(filepath.with_suffix(".html"))
            raise ImportError(
                "ReportLab no instalado. Se genero HTML en su lugar. "
                "Instalar: pip install reportlab"
            )

        # Colores JARVIS
        CYAN = HexColor("#00b4ff")
        DARK_BG = HexColor("#0a1a2a")
        TEXT_COLOR = HexColor("#1a1a2e")
        SUBTITLE_COLOR = HexColor("#5a8aaa")
        ACCENT = HexColor("#0066cc")

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=25 * mm,
            leftMargin=25 * mm,
            topMargin=25 * mm,
            bottomMargin=20 * mm,
            title=title,
            author=author,
        )

        # Estilos
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "JarvisTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=ACCENT,
            spaceAfter=6,
            alignment=TA_LEFT,
        )

        subtitle_style = ParagraphStyle(
            "JarvisSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=11,
            textColor=SUBTITLE_COLOR,
            spaceAfter=4,
        )

        meta_style = ParagraphStyle(
            "JarvisMeta",
            parent=styles["Normal"],
            fontName="Courier",
            fontSize=8,
            textColor=HexColor("#888888"),
            spaceAfter=16,
        )

        body_style = ParagraphStyle(
            "JarvisBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            textColor=TEXT_COLOR,
            leading=16,
            spaceAfter=8,
        )

        section_title_style = ParagraphStyle(
            "JarvisSectionTitle",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            textColor=ACCENT,
            spaceBefore=16,
            spaceAfter=8,
            borderWidth=0,
            borderPadding=0,
            leftIndent=0,
        )

        footer_style = ParagraphStyle(
            "JarvisFooter",
            parent=styles["Normal"],
            fontName="Courier",
            fontSize=7,
            textColor=HexColor("#aaaaaa"),
            alignment=TA_CENTER,
            spaceBefore=20,
        )

        # Construir documento
        story = []

        # Titulo
        story.append(Paragraph(title, title_style))
        if subtitle:
            story.append(Paragraph(subtitle, subtitle_style))

        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")
        story.append(Paragraph(f"Autor: {author}  |  Fecha: {date_str}", meta_style))

        # Linea separadora
        story.append(HRFlowable(
            width="100%", thickness=1, color=CYAN,
            spaceAfter=12, spaceBefore=4,
        ))

        # Contenido principal
        if content:
            for paragraph in content.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    # Escapar caracteres XML
                    safe_p = (paragraph
                              .replace("&", "&amp;")
                              .replace("<", "&lt;")
                              .replace(">", "&gt;")
                              .replace("\n", "<br/>"))
                    story.append(Paragraph(safe_p, body_style))

        # Secciones
        if sections:
            for sec in sections:
                sec_title = sec.get("title", "")
                sec_content = sec.get("content", "")

                if sec_title:
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(sec_title, section_title_style))
                    story.append(HRFlowable(
                        width="40%", thickness=0.5, color=ACCENT,
                        spaceAfter=8,
                    ))

                if sec_content:
                    for paragraph in sec_content.split("\n\n"):
                        paragraph = paragraph.strip()
                        if paragraph:
                            safe_p = (paragraph
                                      .replace("&", "&amp;")
                                      .replace("<", "&lt;")
                                      .replace(">", "&gt;")
                                      .replace("\n", "<br/>"))
                            story.append(Paragraph(safe_p, body_style))

        # Footer
        story.append(Spacer(1, 20))
        story.append(HRFlowable(
            width="60%", thickness=0.5, color=HexColor("#cccccc"),
            spaceAfter=8,
        ))
        story.append(Paragraph(
            f"Generado por {author} // {datetime.now().isoformat()} // GENESIS Document System",
            footer_style,
        ))

        doc.build(story)

    def _render_docx(self, filepath: Path, title: str, content: str,
                     author: str, subtitle: str, sections: list):
        """Genera documento Word DOCX con python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError(
                "python-docx no instalado. Instalar: pip install python-docx"
            )

        doc = Document()

        # Configurar margenes
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Estilo titulo
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 102, 204)
        title_run.font.name = "Calibri"

        # Subtitulo
        if subtitle:
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(12)
            sub_run.font.color.rgb = RGBColor(90, 138, 170)
            sub_run.font.italic = True

        # Metadata
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(
            f"Autor: {author}  |  Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        meta_run.font.size = Pt(8)
        meta_run.font.color.rgb = RGBColor(136, 136, 136)
        meta_run.font.name = "Consolas"

        # Linea separadora
        doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0, 180, 255)

        # Contenido principal
        if content:
            for paragraph in content.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph)
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

        # Secciones
        if sections:
            for sec in sections:
                sec_title = sec.get("title", "")
                sec_content = sec.get("content", "")

                if sec_title:
                    heading = doc.add_heading(sec_title, level=2)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(0, 102, 204)

                if sec_content:
                    for paragraph in sec_content.split("\n\n"):
                        paragraph = paragraph.strip()
                        if paragraph:
                            p = doc.add_paragraph()
                            run = p.add_run(paragraph)
                            run.font.size = Pt(10)
                            run.font.name = "Calibri"

        # Footer
        doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(200, 200, 200)
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Generado por {author} // {datetime.now().isoformat()} // GENESIS Document System"
        )
        footer_run.font.size = Pt(7)
        footer_run.font.color.rgb = RGBColor(170, 170, 170)

        # Metadata del documento
        doc.core_properties.author = author
        doc.core_properties.title = title
        doc.core_properties.comments = "Generado por GENESIS AI"

        doc.save(str(filepath))

    # ================================================================
    # UTILIDADES
    # ================================================================

    @staticmethod
    def _safe_filename(title: str) -> str:
        """Convierte titulo a nombre de archivo seguro."""
        import re
        safe = re.sub(r'[^\w\s-]', '', title)
        safe = re.sub(r'[\s]+', '_', safe)
        return safe[:60].strip("_") or "documento"

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        """Formatea bytes a string legible."""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"

    def list_documents(self, limit: int = 20) -> str:
        """Lista documentos generados."""
        if not self.generated_docs:
            return "  No hay documentos generados aun."

        lines = [f"  === DOCUMENTOS GENERADOS ({len(self.generated_docs)}) ==="]
        for doc in self.generated_docs[-limit:]:
            ts = datetime.fromtimestamp(doc["timestamp"]).strftime("%H:%M:%S")
            lines.append(
                f"  [{ts}] [{doc['format'].upper():4s}] "
                f"{doc['title'][:40]} ({doc['size_human']})"
            )
            lines.append(f"         {doc['path']}")
        return "\n".join(lines)
